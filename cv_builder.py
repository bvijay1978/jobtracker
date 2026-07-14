"""Generate a first-draft CV (.docx) for a role from a local profile.

Reads profile data from a JSON file (kept out of the repo — see
profile.example.json) and slots the role's title into the headline. This is a
fast, consistent first cut so you can apply quickly; deeper JD-tailoring on the
few roles you'll actually submit is a job for the agent layer (ask Claude).
"""

from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

import config

NAVY = RGBColor(0x0D, 0x2B, 0x55)
TEAL = RGBColor(0x00, 0x70, 0x9E)
MID = RGBColor(0x44, 0x44, 0x44)
LIGHT = RGBColor(0x76, 0x76, 0x76)
FONT = "Calibri"


def _set(run, size, color, bold=False, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic


def _divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    for k, v in (("w:val", "single"), ("w:sz", "12"), ("w:space", "1"), ("w:color", "0D2B55")):
        bottom.set(qn(k), v)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    _set(p.add_run(text.upper()), 10, NAVY, bold=True)
    _divider(doc)


def _job_header(doc, title, company, dates):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), "9026")
    tabs.append(tab)
    pPr.append(tabs)
    _set(p.add_run(title), 10.5, NAVY, bold=True)
    _set(p.add_run(" | "), 10.5, MID)
    _set(p.add_run(company), 10.5, TEAL, italic=True)
    _set(p.add_run("\t"), 10, LIGHT)
    _set(p.add_run(dates), 10, LIGHT, italic=True)


def _bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    _set(p.add_run(text), 10, MID)


def _safe(name: str) -> str:
    return re.sub(r'[\\/:*?"<>|]+', "", (name or "")).strip() or "Role"


def load_profile(profile_path: Path | str | None = None) -> dict:
    profile_path = Path(profile_path) if profile_path else config.PROFILE_PATH
    if not profile_path or not Path(profile_path).exists():
        raise FileNotFoundError(
            f"Profile not found at {profile_path}. Copy profile.example.json to "
            "profile.json (or set JOBTRACKER_PROFILE) and fill in your details."
        )
    return json.loads(Path(profile_path).read_text(encoding="utf-8"))


# Detects the role's sector from its title / fit-notes / company so the builder
# can lead with the most relevant profile and competencies — a rule-based,
# offline first cut (deeper JD-tailoring is an agent task; see docs/ARCHITECTURE).
SECTOR_PATTERNS = {
    "fs": r"financ|\bbank|payment|settlement|\bfx\b|insuranc|underwrit|wealth|"
          r"trading|fintech|regulator|\bfca\b|\bfsa\b|capital market|broking",
    "ai": r"\bai\b|artificial intelligence|machine learning|\bml\b|data science|"
          r"\bnhs\b|clinical|governance|\bgenai\b|\bllm\b|cognitive",
    "public": r"public sector|\bgov\b|government|council|civil service|"
              r"local authorit|ministr|department for",
}


def detect_sector(role: dict) -> str:
    text = " ".join(str(role.get(k, "")) for k in ("title", "fit_notes", "company")).lower()
    for sector, pattern in SECTOR_PATTERNS.items():
        if re.search(pattern, text):
            return sector
    return "default"


def resolve_profile(prof: dict, sector: str, max_competencies: int = 10) -> tuple[str, list]:
    """Pick the sector's profile paragraph and order competencies (lead ones first)."""
    profiles = prof.get("profiles")
    if profiles:
        profile_text = profiles.get(sector) or profiles.get("default") or ""
    else:  # legacy single-profile format
        profile_text = prof.get("profile", "")
    base = list(prof.get("competencies", []))
    lead = list((prof.get("leadCompetencies") or {}).get(sector, []))
    ordered = lead + [c for c in base if c not in lead]
    return profile_text, ordered[:max_competencies]


def generate_cv(role: dict, profile_path=None, out_dir=None) -> Path:
    """Write a first-draft CV for ``role`` and return the saved path."""
    prof = load_profile(profile_path)
    out_dir = Path(out_dir) if out_dir else config.CV_DIR
    out_dir.mkdir(parents=True, exist_ok=True)
    title = (role.get("title") or "Role").strip()
    sector = detect_sector(role)
    profile_text, competencies = resolve_profile(prof, sector)
    path = out_dir / f"{prof.get('name', 'CV')} - {_safe(title)}.docx"

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    _set(p.add_run(prof.get("name", "")), 22, NAVY, bold=True)

    p = doc.add_paragraph()  # title line == the role title
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    _set(p.add_run(title), 11, TEAL)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    _set(p.add_run(prof.get("contact", "")), 9.5, LIGHT)
    _divider(doc)

    _heading(doc, "Profile")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    _set(p.add_run(profile_text), 10, MID)

    _heading(doc, "Core Competencies")
    comps = competencies
    half = (len(comps) + 1) // 2
    left, right = comps[:half], comps[half:]
    table = doc.add_table(rows=max(len(left), len(right), 1), cols=2)
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            for edge in ("top", "bottom", "left", "right"):
                e = OxmlElement(f"w:{edge}")
                e.set(qn("w:val"), "none")
                borders.append(e)
            tcPr.append(borders)
    for i, comp in enumerate(left):
        cell = table.cell(i, 0)
        cell.text = ""
        _set(cell.paragraphs[0].add_run(f"▪ {comp}"), 10, MID)
    for i, comp in enumerate(right):
        cell = table.cell(i, 1)
        cell.text = ""
        _set(cell.paragraphs[0].add_run(f"▪ {comp}"), 10, MID)

    _heading(doc, "Professional Experience")
    for job in prof.get("jobs", []):
        _job_header(doc, job.get("title", ""), job.get("company", ""), job.get("dates", ""))
        for bullet in job.get("bullets", []):
            _bullet(doc, bullet)

    _heading(doc, "Key Achievements")
    for ach in prof.get("achievements", []):
        _bullet(doc, ach)

    _heading(doc, "Certifications & Education")
    if prof.get("certifications"):
        _set(doc.add_paragraph().add_run(prof["certifications"]), 10, MID)
    if prof.get("education"):
        _set(doc.add_paragraph().add_run(prof["education"]), 10, MID)

    if prof.get("rightToWork"):
        _heading(doc, "Additional")
        _set(doc.add_paragraph().add_run(prof["rightToWork"]), 10, MID)

    doc.save(str(path))
    return path
