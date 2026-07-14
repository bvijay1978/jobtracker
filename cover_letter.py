"""Generate a draft cover letter (.docx) for a tracked role.

Two tiers, same as the CV split (ADR-007/008):

- **Offline default** (no LLM, what running the in-app button gives you): a
  structured draft seeded with your actual background — the sector-aware
  summary and top competencies from your profile (same source and detection
  as cv_builder) — so the body paragraph and signature are filled in, not
  left as brackets. Not JD-tailored.
- **JD-tailored** (``body_paragraphs``, ADR-012): when a role is queued as
  **Draft CV & Cover Letter**, Claude reads the job description alongside the
  screening CV and writes 1-2 paragraphs mirroring the JD's language against
  your *genuine* experience — same honesty rule as the screening CV, no
  invented claims — and passes them here instead of the offline summary.
"""

from __future__ import annotations

import datetime as dt
import re
from pathlib import Path

from docx import Document

import config
from cv_builder import detect_sector, load_profile, resolve_profile


def _safe(name: str) -> str:
    """Make a string safe to use in a filename."""
    cleaned = re.sub(r'[\\/:*?"<>|]+', "", (name or "")).strip()
    return cleaned or "Untitled"


def cover_letter_path(job: dict, out_dir: Path | str | None = None) -> Path:
    out_dir = Path(out_dir) if out_dir else config.COVER_LETTER_DIR
    company = _safe(job.get("company") or "Company")
    title = _safe(job.get("title") or "Role")
    return out_dir / f"Cover Letter - {company} - {title}.docx"


def generate_cover_letter(
    job: dict,
    out_dir: Path | str | None = None,
    profile_path: Path | str | None = None,
    body_paragraphs: list[str] | None = None,
) -> Path:
    """Write a draft cover letter for ``job`` and return the saved path.

    ``body_paragraphs``, when supplied, are inserted verbatim as the letter's
    body in place of the offline profile summary — this is where Claude's
    JD-tailored paragraphs go (see the module docstring / ADR-012).
    """
    out_dir = Path(out_dir) if out_dir else config.COVER_LETTER_DIR
    out_dir.mkdir(parents=True, exist_ok=True)
    path = cover_letter_path(job, out_dir)

    company = (job.get("company") or "the company").strip()
    title = (job.get("title") or "the role").strip()
    location = (job.get("location") or "").strip()
    fit = (job.get("fit_notes") or "").strip()
    today = dt.date.today().strftime("%d %B %Y")

    try:
        prof = load_profile(profile_path)
    except FileNotFoundError:
        prof = {}
    name = (prof.get("name") or "").strip()
    summary, competencies = ("", [])
    if prof:
        summary, competencies = resolve_profile(prof, detect_sector(job), max_competencies=3)

    doc = Document()
    doc.add_paragraph(today)
    doc.add_paragraph("")
    doc.add_paragraph("Dear Hiring Manager,")

    intro = (
        f"I am writing to apply for the {title} position at {company}"
        + (f" ({location})" if location else "")
        + ". Having reviewed the role, I am confident that my background and "
        "experience make me a strong fit."
    )
    doc.add_paragraph(intro)

    if body_paragraphs:
        # JD-tailored (Claude authored these with the JD in hand) — self-contained,
        # no offline fit-notes sentence appended.
        for para in body_paragraphs:
            doc.add_paragraph(para)
    else:
        if summary:
            body = summary
            if competencies:
                lead = ", ".join(competencies[:-1])
                body += (
                    f" My core strengths include {lead}"
                    f"{' and ' if lead else ''}{competencies[-1]}."
                    if len(competencies) > 1
                    else f" My core strengths include {competencies[0]}."
                )
        else:
            body = (
                "In my career to date, [add two or three sentences on your most relevant "
                "experience, skills and achievements for this role]."
            )
        if fit:
            body += f" This role stood out to me in particular because {fit[0].lower()}{fit[1:]}."
        doc.add_paragraph(body)

    doc.add_paragraph(
        f"I would welcome the opportunity to discuss how I can contribute to {company}. "
        "Thank you for considering my application; I look forward to hearing from you."
    )
    doc.add_paragraph("")
    doc.add_paragraph("Yours sincerely,")
    doc.add_paragraph(name or "[Your name]")

    doc.save(str(path))
    return path
