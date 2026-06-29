import config
import screening_cv

EXAMPLE = config.PROJECT_ROOT / "profile.example.json"


def test_keyword_coverage_matches_terms():
    cov = screening_cv.keyword_coverage(
        ["Scrum", "SAFe", "PMO", "stakeholder management"],
        "Experienced in Scrum and PMO delivery with strong stakeholder management.",
    )
    assert "Scrum" in cov["covered"]
    assert "PMO" in cov["covered"]
    assert "stakeholder management" in cov["covered"]
    assert "SAFe" in cov["missing"]
    assert cov["pct"] == 75


def test_keyword_coverage_word_boundary():
    # "AI" must not match inside "detail"; should match a standalone "AI".
    assert screening_cv.keyword_coverage(["AI"], "rich detail here")["missing"] == ["AI"]
    assert screening_cv.keyword_coverage(["AI"], "AI governance lead")["covered"] == ["AI"]


def test_generate_screening_cv_is_ats_safe(tmp_path):
    role = {"title": "AI Governance Lead"}
    screening = {
        "target_title": "AI Governance Lead",
        "summary": "AI governance and assurance leader in regulated environments.",
        "core_skills": ["AI governance", "Risk & controls", "PMO", "Stakeholder management"],
        "experience": [
            {"title": "Lead", "company": "X", "dates": "2024", "bullets": ["Ran AI governance."]}
        ],
    }
    cv = screening_cv.generate_screening_cv(
        role, screening, profile_path=EXAMPLE, out_dir=tmp_path
    )
    assert cv.exists()

    from docx import Document

    doc = Document(str(cv))
    assert len(doc.tables) == 0  # ATS-safe: no tables/columns to confuse a parser
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "AI Governance Lead" in text  # title mirrors the advert
    assert "AI governance" in text  # keyword block present
